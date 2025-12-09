# app_streamlit.py
import streamlit as st
import pandas as pd
import io
import tempfile
import time
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from googleapiclient.errors import HttpError
import gspread
from docx import Document
from docx.shared import Inches

# ---------------- Config ----------------
# TEMPLATE DOC ID (Google Docs template)
TEMPLATE_DOC_ID = "1rqRTFLG3D0v4Sz7mnJ9iREOxyl1YWLexXyJVMJa_Ezk"  # giữ như trong code cũ (cần quyền)
SPREADSHEET_URL = "https://docs.google.com/spreadsheets/d/1lG4NN4kdk-bcKqWh0fYYRorMSCIypWqsByJvfwO8KgY/edit"

# ------------- Utils: creds -------------
@st.cache_resource
def get_credentials():
    # Load service account info from st.secrets
    sa_info = st.secrets["gcp_service_account"]
    creds = service_account.Credentials.from_service_account_info(
        sa_info,
        scopes=[
            "https://www.googleapis.com/auth/drive",
            "https://www.googleapis.com/auth/documents",
            "https://www.googleapis.com/auth/spreadsheets",
        ],
    )
    return creds

@st.cache_resource
def get_api_clients():
    creds = get_credentials()
    docs_service = build("docs", "v1", credentials=creds)
    drive_service = build("drive", "v3", credentials=creds)
    sheets_client = gspread.authorize(creds)
    return docs_service, drive_service, sheets_client

# ------------- Read sheets (CSDL + Taichinh) -------------
@st.cache_data(ttl=300)
def load_data():
    _, _, sheets_client = get_api_clients()
    # open by url
    sh = sheets_client.open_by_url(SPREADSHEET_URL)
    try:
        sheet_csdl = sh.worksheet("CSDL")
        sheet_taichinh = sh.worksheet("Taichinh")
    except Exception as e:
        st.error(f"Không tìm thấy sheet CSDL hoặc Taichinh: {e}")
        return pd.DataFrame(), pd.DataFrame()

    data_csdl = pd.DataFrame(sheet_csdl.get_all_records())
    data_taichinh = pd.DataFrame(sheet_taichinh.get_all_records())

    # Strip column names
    data_csdl.columns = [c.strip() for c in data_csdl.columns]
    data_taichinh.columns = [c.strip() for c in data_taichinh.columns]

    # months list from first column of Taichinh (exclude header)
    thang_vals = [v for v in sheet_taichinh.col_values(1) if v and v.lower() != "thang"]
    thang_list = sorted(set(thang_vals), key=str)

    return data_csdl, data_taichinh, thang_list

# ------------- Create Google Doc (copy template + replace placeholders) -------------
def create_google_doc_copy_and_replace(user_data: dict):
    docs_service, drive_service, _ = get_api_clients()
    # copy template
    title = f"BBNT_{user_data.get('ma_tram', 'Unknown')}_{user_data.get('Thang', '')}_{int(time.time())}"
   
    TARGET_FOLDER_ID = "1k9isJyWgX3fq2Lh_FEIDwvMPbYxvI3bW"

    copied = drive_service.files().copy(
        fileId=TEMPLATE_DOC_ID,
        body={
            "name": title,
            "parents": [TARGET_FOLDER_ID]
        },
        supportsAllDrives=True
    ).execute()

    new_doc_id = copied.get("id")

    # compute derived fields (same logic as original)
    user_data = dict(user_data)  # copy
    user_data["Danh_gia_cot"] = "Đạt" if user_data.get("Loai_cot") == "cột dây co" else "Không đánh giá"
    user_data["Danh_gia_PM"] = "Đạt" if user_data.get("Phong_may") != "Không thuê" else "Không đánh giá"
    user_data["Danh_gia_DH"] = "Đạt" if user_data.get("Dieu_hoa") != "Không thuê" else "Không đánh giá"

    try:
        tong_tien = user_data.get("tongtienky", 0)
        # may be string -> try convert
        try:
            tong_val = float(tong_tien)
            # convert to int if it's integral
            if tong_val.is_integer():
                tong_val = int(tong_val)
            user_data["Tien_bang_chu"] = str(tong_val)
        except Exception:
            user_data["Tien_bang_chu"] = str(tong_tien)
    except Exception:
        user_data["Tien_bang_chu"] = ""

    # build replace requests
    requests = []
    for key, val in user_data.items():
        requests.append({
            "replaceAllText": {
                "containsText": {"text": f"${key}", "matchCase": True},
                "replaceText": str(val)
            }
        })

    docs_service.documents().batchUpdate(documentId=new_doc_id, body={"requests": requests}).execute()
    return new_doc_id

# ------------- Export Google Doc to DOCX bytes -------------
def export_docx_bytes(doc_id):
    _, drive_service, _ = get_api_clients()
    try:
        request = drive_service.files().export_media(
            fileId=doc_id,
            mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        fh.seek(0)
        return fh.read()
    except HttpError as e:
        st.error(f"Google Drive export error: {e}")
        return None

# ------------- Insert images into docx bytes and return new bytes -------------
def insert_images_into_docx_bytes(docx_bytes: bytes, images_map: dict):
    # images_map: {placeholder: uploaded_file (streamlit UploadedFile)}
    doc = Document(io.BytesIO(docx_bytes))

    # helper to insert into paragraph or table cell where placeholder occurs
    def insert_picture_in_paragraph_or_cell(placeholder, image_path):
        # paragraphs
        for p in doc.paragraphs:
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, "")
                run = p.add_run()
                run.add_picture(image_path, width=Inches(3))
                return True
        # tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, "")
                        run = cell.paragraphs[0].add_run()
                        run.add_picture(image_path, width=Inches(3))
                        return True
        return False

    # Save each uploaded image to temp file and insert
    for placeholder, upfile in images_map.items():
        if upfile is None:
            continue
        # write to temp file
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
        tmp.write(upfile.getbuffer())
        tmp.flush()
        tmp.close()
        inserted = insert_picture_in_paragraph_or_cell(placeholder, tmp.name)
        # if not inserted, just append image at end of document
        if not inserted:
            doc.add_paragraph("")
            run = doc.paragraphs[-1].add_run()
            run.add_picture(tmp.name, width=Inches(3))

    # return bytes
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()

# ------------- Delete file on Drive -------------
def delete_drive_file(file_id):
    _, drive_service, _ = get_api_clients()
    try:
        drive_service.files().delete(fileId=file_id).execute()
    except Exception:
        pass

# ------------- Streamlit UI -------------
st.set_page_config(page_title="BBNT - XH Hóa (Streamlit)", layout="centered")
st.title("BBNT - Xã Hội Hóa (Streamlit)")

st.info("Ứng dụng chạy bằng Service Account lưu trong st.secrets. Không upload file JSON lên repo.")

data_csdl, data_taichinh, thang_list = load_data()

if data_csdl.empty:
    st.error("Dữ liệu CSDL rỗng — kiểm tra share sheet và quyền Service Account.")
    st.stop()

# Prepare lists
csdl_dict = data_csdl.to_dict(orient="list")
ma_tram_list = [str(x).strip().upper() for x in csdl_dict.get("ma_tram", [])]
Password_list = csdl_dict.get("Password", [])

with st.form("login_form"):
    ma_tram_input = st.text_input("Mã Nhà Trạm (ví dụ: ABC123)").strip().upper()
    password_input = st.text_input("Mật khẩu", type="password")
    thang_chon = st.selectbox("Tháng thanh toán", options=thang_list) if thang_list else st.text_input("Tháng")
    submitted = st.form_submit_button("Đăng nhập & Tạo Docs")

if submitted:
    if not ma_tram_input or not password_input or not thang_chon:
        st.warning("Nhập đầy đủ thông tin.")
    else:
        if ma_tram_input in ma_tram_list:
            idx = ma_tram_list.index(ma_tram_input)
            correct_pw = Password_list[idx]
            if password_input == str(correct_pw):
                st.success("Đăng nhập thành công!")
                # build user_data
                user_data = {col: csdl_dict[col][idx] for col in csdl_dict.keys()}
                # find financial row
                try:
                    # ensure types string for comparison
                    df = data_taichinh.copy()
                    df["Ma_vi_tri"] = df["Ma_vi_tri"].astype(str).str.upper()
                    df["Thang"] = df["Thang"].astype(str)
                    match = df[(df["Ma_vi_tri"] == ma_tram_input) & (df["Thang"] == str(thang_chon))]
                except Exception:
                    match = pd.DataFrame()

                if match.empty:
                    st.error(f"Không tìm thấy dữ liệu thanh toán cho tháng {thang_chon}.")
                else:
                    # merge first match
                    user_data.update(match.iloc[0].to_dict())
                    user_data["Thang"] = thang_chon
                    st.info("Đang tạo Google Docs trên Drive...")
                    try:
                        doc_id = create_google_doc_copy_and_replace(user_data)
                        st.success("Đã tạo Google Docs.")
                        doc_url = f"https://docs.google.com/document/d/{doc_id}/edit"
                        st.markdown(f"- [Mở Google Doc đã tạo]({doc_url})")
                        # Export to docx
                        st.info("Đang xuất DOCX từ Google Docs...")
                        docx_bytes = export_docx_bytes(doc_id)
                        if not docx_bytes:
                            st.error("Không thể export DOCX. Kiểm tra quyền Drive/Docs của Service Account.")
                        else:
                            st.success("Đã export DOCX. Bây giờ upload ảnh (nếu cần) để chèn vào file.")
                            # Ask user to upload up to 8 images
                            st.write("Upload tối đa 8 ảnh (theo yêu cầu). Bỏ trống nếu không cần chèn.")
                            placeholders = [
                                "${Anh1}", "${Anh2}", "${Anh3}", "${Anh4}",
                                "${Anh5}", "${Anh6}", "${Anh7}", "${Anh8}"
                            ]
                            uploaded_images = {}
                            cols = st.columns(4)
                            for i, ph in enumerate(placeholders):
                                with cols[i % 4]:
                                    uploaded_images[ph] = st.file_uploader(f"{ph}", type=["jpg","jpeg","png"], key=f"img{i}")

                            if st.button("Chèn ảnh và Tải file DOCX"):
                                st.info("Đang chèn ảnh vào file DOCX...")
                                result_bytes = insert_images_into_docx_bytes(docx_bytes, uploaded_images)
                                if result_bytes:
                                    fname = f"BBNT_{ma_tram_input}_{thang_chon}.docx"
                                    st.download_button("⬇️ Tải file DOCX cuối cùng", data=result_bytes, file_name=fname, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                                    st.success("Hoàn tất! File đã sẵn sàng để tải.")
                                else:
                                    st.error("Lỗi khi chèn ảnh.")
                            # Option to download original (no images)
                            st.download_button("Tải DOCX (không chèn ảnh)", data=docx_bytes, file_name=f"BBNT_{ma_tram_input}_{thang_chon}_raw.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                        # Delete doc on Drive (same as original behavior)
                        if st.button("Xóa tệp Google Docs trên Drive (không bắt buộc)"):
                            delete_drive_file(doc_id)
                            st.info("Đã gửi lệnh xóa (nếu file tồn tại và có quyền).")
                    except Exception as e:
                        st.exception(e)
            else:
                st.error("Mật khẩu không đúng.")
        else:
            st.error("Mã nhà trạm không tồn tại trong CSDL.")
